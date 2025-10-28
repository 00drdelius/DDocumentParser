import os
import asyncio
import aiohttp
import subprocess
import asyncio.subprocess as asubprocess
from typing import *
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.oxml.ns import qn

from .logg import logger
from .config import MINERU_URL, HTTP_CLIENT

T = TypeVar("T")

async def async_wrapper(callable: Callable[..., T], *args, **kwargs)-> Optional[T]:
    """wrap sync function to be async"""
    result =  await asyncio.to_thread(callable, *args, **kwargs)
    return result


def check_libreoffice():
    try:
        #NOTE set `check=True` to raise exception if command fails
        process = subprocess.run(["soffice", "--version"],shell=True, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except (subprocess.CalledProcessError, FileNotFoundError):
        raise OSError("LibreOffice is not installed or unavailable. Please install LibreOffice and ensure it is in the system PATH.")
    else:
        if process.returncode!=0:
            raise OSError("LibreOffice is not installed or unavailable. Please install LibreOffice and ensure it is in the system PATH.")


async def acheck_libreoffice():
    try:
        process = await asubprocess.create_subprocess_exec(*["soffice", "--version"], stdout=asubprocess.PIPE, stderr=asubprocess.PIPE)
        stdout, stderr = await process.communicate()
    except (FileNotFoundError, subprocess.CalledProcessError):
        raise OSError("LibreOffice is not installed or unavailable. Please install LibreOffice and ensure it is in the system PATH.")
    else:
        if process.returncode!=0:
            raise OSError("LibreOffice is not installed or unavailable. Please install LibreOffice and ensure it is in the system PATH.")


def convert_docs_to_docxs(
    input_directory_or_file:Union[str, Path], output_directory:Union[str,Path]=None) -> list[Path]:
    """
    single convert or batch convert all .doc files in the input_directory_or_file (including subdirectories) to .docx format
    using LibreOffice's command line interface.
    Args:
        input_directory_or_file (str | Path): input directory containing .doc files or filepath to a single .doc file
        output_directory (str | Path): output directory path where converted .docx files will be saved
    Raises:
        ValueError: If the input directory does not exist or is not a directory
        OSError: If LibreOffice is not installed or unavailable
    Returns:
        out(list[Path]): list of converted file paths
    """


    input_path = Path(input_directory_or_file) if not isinstance(input_directory_or_file,Path) else input_directory_or_file
    output_dir = Path(output_directory) if not isinstance(output_directory, Path) else output_directory

    if not input_path.exists():
        raise ValueError(f"Input path {input_path} does not exist.")

    # Create output directory if it does not exist
    output_dir.mkdir(parents=True, exist_ok=True)
    if input_path.is_file():
        file_option = str(input_path)
    elif input_path.is_dir():
        file_option = str(f'{str(input_path)}/**/*.doc')

    command = f'soffice --headless --convert-to docx --outdir "{str(output_dir)}" "{file_option}"'
    try:
        logger.info("output_dir is ", str(output_dir))
        logger.info("start converting by LibreOffice...")
        result=subprocess.run(command, shell=True, check=True,stderr=subprocess.PIPE, text=True, stdout=subprocess.PIPE)
        logger.info("LibreOffice output:", result.stdout or result.stderr)
        if result.returncode == 0:
            logger.info("Conversion successful")
        else:
            raise OSError(f"Error occurred during conversion: {result.stderr.strip()}")
    except subprocess.CalledProcessError as e:
        raise OSError(f"Error occurred during conversion: {e.stderr.strip()}") from e
    else:
        docx_filepaths = list(output_dir.glob("*.docx"))
    return docx_filepaths


async def aconvert_docs_to_docxs(
    input_directory_or_file:Union[str, Path], output_directory:Union[str,Path]=None) -> list[Path]:
    """
    **async version**

    single convert or batch convert all .doc files in the input_directory_or_file (including subdirectories) to .docx format
    using LibreOffice's command line interface.
    Args:
        input_directory_or_file (str | Path): input directory containing .doc files or filepath to a single .doc file
        output_directory (str | Path): output directory path where converted .docx files will be saved
    Raises:
        ValueError: If the input directory does not exist or is not a directory
        OSError: If LibreOffice is not installed or unavailable
    """
    input_path = Path(input_directory_or_file) if not isinstance(input_directory_or_file,Path) else input_directory_or_file
    output_dir = Path(output_directory) if not isinstance(output_directory, Path) else output_directory

    if not input_path.exists():
        raise ValueError(f"Input path {input_path} does not exist.")

    # Create output directory if it does not exist
    output_dir.mkdir(parents=True, exist_ok=True)
    if input_path.is_file():
        file_option = str(input_path)
    elif input_path.is_dir():
        file_option = str(f'{str(input_path)}/**/*.doc')

    command = [
        'soffice','--headless','--convert-to docx','--outdir',output_dir, file_option ]
    try:
        logger.info("output_dir is ", str(output_dir))
        logger.info("start converting by LibreOffice...")
        process = await asubprocess.create_subprocess_exec(*command, stdout=asubprocess.PIPE, stderr=asubprocess.PIPE)
        stdout, stderr = await process.communicate()
        stdout = stdout.decode("utf8")
        stderr = stderr.decode("utf8")
        logger.info("LibreOffice output:", stdout or stderr)
        if process.returncode == 0:
            logger.info("Conversion successful")
        else:
            raise OSError(f"Error occurred during conversion: {stderr}")
    except subprocess.CalledProcessError as e:
        raise OSError(f"Error occurred during conversion: {stderr}") from e
    else:
        docx_filepaths = list(output_dir.glob("*.docx"))
    return docx_filepaths

def get_pure_pdf_text(
    file:str | Path | bytes,
    exclude_header:bool = False,
    exclude_footer:bool = False,
    exclude_pixels:int = 60,
    ) -> list[str]:
    """
    extract pure text from a given PDF file, with header or footer removed. 

    **[SPECIAL ADDRESS]** It excludes text in the header and footer areas (top and bottom 60pt of each page),
    which often contain page numbers, document titles, or other repetitive information that may interfere with main text.

    Args:
        file(str| Path | bytes): PDF filepath or PDF file bytes
        exclude_header(bool): given True to exclude header
        exclude_footer(bool): given True to exclude footer
        exclude_pt(int): how many pt you want to exclude in header or footer?\
        (See refer to https://en.wikipedia.org/wiki/Point_(typography) for more details on **pt** unit)
    Returns:
        out(list[str]): list of pure texts extracted from all pdf pages.
    """
    full_texts=[]
    if isinstance(file, (str, Path,)):
        pdf_doc = fitz.open(str(file))
    elif isinstance(file, bytes):
        pdf_doc = fitz.open(stream=file, filetype="pdf")
    for page in pdf_doc:
        rect = page.rect
        header_area=None
        footer_area=None
        if exclude_header:
            header_area = fitz.Rect(rect.x0, rect.y0, rect.x1, rect.y0 + exclude_pixels)  # usually top 60pt, header
        if exclude_footer:
            footer_area = fitz.Rect(rect.x0, rect.y1 - exclude_pixels, rect.x1, rect.y1)  # usually bottom 60pt, footer

        # extract main text（exclude header/footer）
        page_text=[]
        words = page.get_text("words")  # (x0, y0, x1, y1, word, block_no, line_no, word_no)
        for w in words:
            if header_area and not header_area.intersects(fitz.Rect(w[:4])):
                page_text.append(w[4])

            if footer_area and not footer_area.intersects(fitz.Rect(w[:4])):
                page_text.append(w[4])
        # main_text = "\n".join(
        #     w[4] for w in words if not header_area.intersects(fitz.Rect(w[:4]))
        #                         and not footer_area.intersects(fitz.Rect(w[:4]))
        # )
        page_text="\n".join(page_text)
        full_texts.append(page_text.strip())

    return full_texts


digit_to_chn_digit={
    1:"一",2:"二",3:"三",4:"四",5:"五",6:"六",7:"七",8:"八",9:"九",10:"十",0:"零"
}

def get_pure_docx_text(filepath: Union[str, Path]) -> str:
    """
    extract pure text from a given .docx file, including **auto numbered list items**,
    which cannot be extracted by simply reading the paragraph text.

    Args:
        file_path (str): .docx filepath
    Returns:
        str: pure text extracted, with line breaks between paragraphs 
    Raises:
        ValueError: If the file path is not a valid .docx file
    """

    doc = Document(filepath)

    try:
        #NOTE address auto numbered list items
        numbering_part = doc.part.numbering_part._element
    except BaseException as e:
        print(f"[Warning] Failed to access numbering part in {filepath}, extracting plain text only.")
        return '\n'.join([para.text for para in doc.paragraphs if para.text])

    ##NOTE get numId to abstractNumId mapping, abstractNumId is the key to auto numbering format
    numId2abstractId = {
        num.numId: num.abstractNumId.val
        for num in numbering_part.num_lst
    }
    ##NOTE get abstractNumId and ilvl to style mapping.
    ###NOTE You may get raw numbering_part by converting docx to zip and `word/numbering.xml` is the file you want.
    ####NOTE see https://learn.microsoft.com/zh-cn/previous-versions/office/ee922775%28v=office.14%29#%E6%A6%82%E8%BF%B0
    ####NOTE see also https://blog.51cto.com/u_11866025/11202906
    abstractNumId2style = {}
    for abstractNumIdTag in numbering_part.findall(qn("w:abstractNum")):
        abstractNumId = abstractNumIdTag.get(qn("w:abstractNumId"))
        for lvlTag in abstractNumIdTag.findall(qn("w:lvl")):
            ilvl = lvlTag.get(qn("w:ilvl"))
            style = {tag.tag[tag.tag.rfind("}") + 1:]: tag.get(qn("w:val"))
                    for tag in lvlTag.xpath("./*[@w:val]", namespaces=numbering_part.nsmap)}
            abstractNumId2style[(int(abstractNumId), int(ilvl))] = style
    
    #NOTE extract text, including auto numbered list items
    abstract_count=dict()  # key: (numId, ilvl), value: count
    full_text=[]
    for paragraph in doc.paragraphs:
        prefix_text=""
        numpr = paragraph._element.pPr.numPr
        if numpr is not None and numpr.numId.val != 0:
            numId = numpr.numId.val
            ilvl = numpr.ilvl.val
            abstractId = numId2abstractId[numId]
            style = abstractNumId2style[(abstractId, ilvl)]
            start_index = int(style.get("start", "1"))
            abstract_count[(numId, ilvl)] = abstract_count.get((numId, ilvl), start_index)
            chn_digit = digit_to_chn_digit[abstract_count[(numId, ilvl)]]
            prefix_text=style.get("lvlText", "").replace("%1", chn_digit)

            abstract_count[(numId, ilvl)] += 1 #NOTE increase index for next item

        if paragraph.text!=None:
            text = prefix_text + " " + paragraph.text.strip()
            full_text.append(text)

    return '\n'.join(full_text)


def get_pure_text(filepath: str) -> str:
    """
    extract pure text from a given .docx or .pdf file.

    Args:
        file_path (str): .docx or .pdf filepath
    Returns:
        str: pure text extracted, with line breaks between paragraphs
    """
    match filepath.lower().rsplit(".", 1)[-1]:
        case "docx":
            return get_pure_docx_text(filepath)
        case "pdf":
            return get_pure_pdf_text(filepath)
        case _:
            raise ValueError(f"Unsupported file type: {filepath}")


async def request_mineru(
    request_id: str,
    output_format: Literal["json", "markdown"],
    file_stream:bytes,
    filename:str,
):
    formdata = aiohttp.FormData()
    formdata.add_field("file",file_stream,filename=filename,)
    formdata.add_field("request_id",request_id)
    formdata.add_field("output_format", output_format)
    async with HTTP_CLIENT.post(MINERU_URL, data=formdata) as aresp:
        aresp.raise_for_status()
        data= await aresp.json()
        return data


if __name__ == "__main__":
    # text=get_pure_docx_text(r"D:\workspaces\ChinaMobile\审计局\审计规章制度(docx)\1.中山市审计局财务管理制度.docx")
    text=get_pure_pdf_text(r"审计规章制度\关于印发《中山市审计局审计业务电子数据管理办法》的通知_已签章_V0(1).pdf")
    with open("test.txt","w",encoding="utf-8") as f:
        f.write(text)